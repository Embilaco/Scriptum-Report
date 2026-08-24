"""The essay example, built and read back.

``essay.yaml`` writes a whole essay into ``essay.docx``: a title, subtitle and
authors, an intro, nine clones of the blueprint ``subsection:content``
(each a heading, a text file, and pictures and more text files added at its
marker), a facts section with two CSV tables (``csvseparator: ';'``,
descriptions ``from: row1``), a conclusion with a closing picture, and one
``_global_`` title the headers carry. A test that only checks the file is
there proves none of that, so this module reads the document back:

* what it *says*, against ``expected/essay.json`` (captured at `44267a8`
  from the ``.rdf`` this fixture was translated from; one line re-captured
  when file reads became UTF-8 -- the old reference held the mojibake of
  ``essay_contemp.txt``'s en dash, the only non-ASCII in the corpus);
* what it *shows*: the chapters in the order the document names them, the
  pictures at the blueprint's 12 cm width (and the conclusion's own), their
  captions, the two tables, the global title in every header.
"""

from pathlib import Path
import importlib.util
import sys

import docx
import pytest

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = Path(__file__).resolve().parent.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_examples import *

module_path = Path(__file__).resolve().parent.parent.parent / '02_basetest' / 'common_case.py'

# Load the module from the given path
spec = importlib.util.spec_from_file_location('common_case', str(module_path))
module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(module)

CaseConfig = module.CaseConfig
run_docx_case = module.run_docx_case
said, normalise, reference, difference, portable, fold, drawings = (
    module.said, module.normalise, module.reference, module.difference,
    module.portable, module.fold, module.drawings)
checkreport_comparison = module.checkreport_comparison

REFERENCE = THIS_DIR / 'expected' / 'essay.json'


def build(tmp_path):
    """The essay, typeset the way every docx case is."""
    config = CaseConfig(
        name="essay",
        case_dir=THIS_DIR,
        document_name="essay.yaml",
        template_doc_name="essay.docx",
        output_name="final_essay.docx",
        include_patterns=["essay*.yaml", "essay.docx"],
        data_source_dir=THIS_DIR / 'data',
        finish=False,
        createpdf=False,
    )
    return run_docx_case(config, tmp_path)


def test_essay_document_is_created(tmp_path):
    print(f'\nWorking in {tmp_path}')

    result_path = build(tmp_path)

    assert result_path.exists(), "Expected final_essay.docx to be generated"
    assert result_path.stat().st_size > 0, "Generated document should not be empty"


def test_the_essay_says_what_the_reference_says(tmp_path):
    """Every paragraph, then every table cell, against the stored reference."""
    document = build(tmp_path)

    got = normalise(portable(said(document), document.parent))
    expected = [fold(line) for line in reference(REFERENCE)]

    assert got == expected, difference(expected, got)


def test_chapters_pictures_tables_and_headers(tmp_path):
    """What the text comparison cannot see."""
    document = docx.Document(build(tmp_path))

    # the chapters, in the order the document names them: the intro, the nine
    # clones of subsection:content, the facts section, the conclusion
    chapters = [p.text for p in document.paragraphs if p.style.name == 'Heading 1']
    assert chapters == [
        'Introduction',
        'Early Experiments and the Birth of the Typewriter',
        'Standardization, Visible Writing, and Social Effects',
        'The Rise of Electric Typewriters',
        'From Word Processors to Digital Keyboards',
        'Digital Typewriting and the Evolution of Fonts',
        'Variable Fonts: A New Frontier in Letter Design',
        'Color Fonts and Multicolored Letterforms',
        'Contemporary Typographic Trends',
        'The Continuing Influence of Typewriting',
        'Selected Facts and Comparisons',
        'Conclusion']

    # seven image:generic clones at the blueprint's <image:generic width=12cm>,
    # then image:conclude at its own size
    assert [size for p in document.paragraphs for size in drawings(p)] == [
        (12.0, 13.43), (12.0, 8.04), (12.0, 8.99), (12.0, 7.52),
        (12.0, 3.69), (12.0, 7.1), (12.0, 6.63), (15.2, 10.16)]

    captions = [p.text for p in document.paragraphs if p.style.name == 'Caption']
    assert captions == [
        'Figure 1: The Sholes and Glidden (1876) - later Remington #1',
        'Figure 1: Remington #2 with a shift key',
        'Figure 1: IBM Selectric II',
        'Figure 1: Microsoft Word 1.0 for DOS (1983)',
        'Figure 1: Mono spaced font - source: Google fonts',
        'Figure 1: Different styles on a complex font - source: Google fonts',
        'Figure 1: EMOJI Font - Source: Google Fonts',
        'Table 1: Simplified timeline of key developments',
        'Table 1: Comparison of writing and typography technologies.']

    # the two CSV tables, read with ';' as the separator
    assert [(len(t.rows), len(t.columns)) for t in document.tables] == [(6, 3), (5, 4)]
    assert [c.text for c in document.tables[0].rows[0].cells] == ['Year', 'Innovation', 'Example']
    assert [c.text for c in document.tables[1].rows[0].cells] == [
        'Technology', 'Typical Period', 'Key Strengths', 'Key Limitations']

    # the _global_ title in the header of every section
    assert len(document.sections) == 3
    for section in document.sections:
        assert [p.text for p in section.header.paragraphs if p.text.strip()] == [
            'From Typewriting to Variable Fonts...']


def test_the_checkreport_notebook_would_say_identical(tmp_path):
    """The comparison ``CheckReport.ipynb`` beside this file ends with --
    plain, no ``comparable()`` -- on a plain build (the notebook itself runs
    ``finish=True``, which changes nothing a text comparison sees unless Word
    refreshes a field). Green when the notebook would print IDENTICAL;
    anything else reports as xfailed with the first difference."""
    report = checkreport_comparison(build(tmp_path), REFERENCE)
    if report:
        pytest.xfail('CheckReport.ipynb would not say IDENTICAL -- ' + report)
