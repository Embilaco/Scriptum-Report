"""The images docx case, built and read back.

``word_images.yaml`` is all pictures: direct fills (``image:main``,
``image:footer``, ``image:inline``, a block with a description), clones of
blueprint image blocks added at a marker (``image:generic``, ``image:lonely``,
``image:small``), an ``icon``, two deliberately missing files, and two
``_global_`` pictures the template tags in both body sections, the header and
the footer. A test that only checks the file is there proves none of that,
so this module reads the document back:

* what it *says*, against ``expected/word_images.json`` (captured at
  `44267a8` from the ``.rdf`` this fixture was translated from, re-captured
  when the user rewrote the fixture, `5bf0c37`/`f9a9789`);
* what it *shows* -- the drawings with their sizes -- which a comparison of
  paragraph texts cannot see: ``ALLOVER `` reads the same with or without a
  picture after it, which is how the global image fill placed nothing from
  `86fc1ef` to `1f75367` while every text reference agreed.
"""

from pathlib import Path
import shutil
import sys

import docx
import pytest

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = Path(__file__).resolve().parent.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *
from common_case import CaseConfig, run_docx_case
from common_case import said, normalise, reference, difference, portable, fold, drawings
from common_case import checkreport_comparison

REFERENCE = THIS_DIR / 'expected' / 'word_images.json'


def build(tmp_path, template='template.docx', case_dir=THIS_DIR):
    """The images case, typeset with *template* from *case_dir*."""
    config = CaseConfig(
        name="report",
        case_dir=case_dir,
        document_name="word_images.yaml",
        template_doc_name=template,
        output_name="final_report.docx",
        include_patterns=["*.yaml", template],
        data_source_dir=DATA_SOURCE,
        finish=False,
        createpdf=False,
    )
    return run_docx_case(config, tmp_path)


def test_document_is_created(tmp_path):
    print(f'\nWorking in {tmp_path}')

    result_path = build(tmp_path)

    assert result_path.exists(), "Expected final_report.docx to be generated"
    assert result_path.stat().st_size > 0, "Generated document should not be empty"


def test_the_document_says_what_the_reference_says(tmp_path):
    """Every paragraph, then every table cell, against the stored reference.
    The two missing files are announced with their path, which the runner
    makes absolute -- portable() takes the workspace out again."""
    document = build(tmp_path)

    got = normalise(portable(said(document), document.parent))
    expected = [fold(line) for line in reference(REFERENCE)]

    assert got == expected, difference(expected, got)


def test_every_picture_is_placed_at_its_size(tmp_path):
    """All eleven inline pictures of the body, in document order, at the size
    the tag or the fill asked for (``width=5cm``, ``height: 8mm``, a
    ``scale``, or the file's own size)."""
    document = docx.Document(build(tmp_path))

    placed = [size for p in document.paragraphs for size in drawings(p)]
    assert placed == [
        (5.0, 5.14),    # image:main        <image:main width=5cm/>, pudding.jpg
        (0.8, 0.8),     # image:allover     global, height 8mm, in section one
        (0.8, 0.8),     # image:allover     global, at the top of section two
        (4.0, 4.0),     # image:inline      <image:inline width=4cm/>, camera.png
        (0.8, 0.8),     # image:allover     global, ALLOVER in section two
        (5.0, 5.0),     # image:generic::1  width=5cm, screw.png
        (5.0, 3.33),    # image:generic::2  width=5cm, curve_test.png
        (4.0, 4.0),     # image:lonely::1   height=4cm, screw.png
        (2.57, 4.02),   # icon:rocket       the template's own picture, copied
        (2.0, 1.33),    # image:small::1    width=2cm, curve_test.png
        (2.0, 2.0),     # image:small::2    width=2cm, screw.png
    ]
    assert len(document.inline_shapes) == 11
    announced = [p.text for p in document.paragraphs if 'no file with image' in p.text]
    assert len(announced) == 2, 'image:existing and image:lonely::2 name not-existing.png'


def test_global_image_fills_place_their_pictures(tmp_path):
    """``_global_`` holds ``image:allover`` (8 mm high) and ``image:header``
    (1 cm). The template tags the first in both body sections and in the
    footer, the second in the header, next to direct fills that were always
    fine. Pinned on the drawings, not the text: the global image branch once
    compared the element's canonical address with the puretag it was asked
    for, matched nothing, and every one of these paragraphs stayed empty --
    which a comparison of paragraph texts cannot see.
    """
    document = docx.Document(build(tmp_path))

    allover = [p for p in document.paragraphs if p.text.startswith('ALLOVER')]
    assert len(allover) == 2, 'one ALLOVER paragraph per body section'
    assert [drawings(p) for p in allover] == [[(0.8, 0.8)], [(0.8, 0.8)]]

    first = document.sections[0]
    header, footer = first.header.paragraphs, first.footer.paragraphs
    assert drawings(header[0]) == [(1.0, 1.0)], '<image:header width=1cm/> at 1 cm high'
    assert header[1].text == 'Here comes the header'
    assert drawings(footer[0]) == [(1.0, 1.5)], 'the direct image:footer fill, 1.5 cm high'
    assert footer[1].text.startswith('Here comes the footer')
    assert drawings(footer[1]) == [(0.8, 0.8)], 'image:allover in the footer'

    left_over = [p.text for p in document.paragraphs + header + footer if '<image:' in p.text]
    assert not left_over


def test_a_header_shared_by_linked_sections_gets_its_picture_once(tmp_path):
    """A section whose header and footer are linked to the previous one reads
    the same paragraphs, so a global fill meets each of them once per section.
    The first visit places the picture and consumes the tag; the second finds
    no tag and has to leave the paragraph alone instead of failing on the run
    it did not get.
    """
    case = tmp_path / 'case'
    case.mkdir()
    shutil.copy(THIS_DIR / 'word_images.yaml', case)
    template = docx.Document(THIS_DIR / 'template.docx')
    template.sections[1].header.is_linked_to_previous = True
    template.sections[1].footer.is_linked_to_previous = True
    template.save(case / 'linked.docx')

    document = docx.Document(build(tmp_path, 'linked.docx', case))

    assert all(s.header.is_linked_to_previous for s in document.sections[1:])
    first = document.sections[0]
    assert drawings(first.header.paragraphs[0]) == [(1.0, 1.0)]
    assert drawings(first.footer.paragraphs[1]) == [(0.8, 0.8)]
    allover = [p for p in document.paragraphs if p.text.startswith('ALLOVER')]
    assert [drawings(p) for p in allover] == [[(0.8, 0.8)], [(0.8, 0.8)]]


def test_the_checkreport_notebook_would_say_identical(tmp_path):
    """The comparison ``CheckReport.ipynb`` beside this file ends with --
    plain, no ``comparable()`` -- on a plain build (the notebook itself runs
    ``finish=True``, which changes nothing a text comparison sees unless Word
    refreshes a field). Green when the notebook would print IDENTICAL;
    anything else reports as xfailed with the first difference."""
    report = checkreport_comparison(build(tmp_path), REFERENCE)
    if report:
        pytest.xfail('CheckReport.ipynb would not say IDENTICAL -- ' + report)
