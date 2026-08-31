"""``breakbefore``: which instances get a page break.

The argument puts a page break in front of a block — **in front of every
instance except the first**. Instance 1 starts wherever the blueprint stood
and needs no break to get there; the ones that follow it do.

The rule changed on 2026-08-31. It used to fire on ``::1`` too, and the essay
showed why that was wrong: its template carries **two manual page breaks** in
front of its blueprint, from before the rule was ever uniform, so its first
content section came out behind *three* of them. Nobody had asked for the
third. "Each of these starts on a new page" is about the ones that follow —
the first one starts where it is.

What the change moved, measured
-------------------------------
Three shipped templates flag blocks ``breakbefore``: ``docx_basic/simple``,
``04_examples/wordreport`` (``tool``, ``preparation``, ``testplan``) and
``04_examples/essay`` (``content``).

* the essay's content sections went from ``3, 1, 1, 1, ...`` to
  ``2, 1, 1, 1, ...`` — the unasked-for third break in front of the first
  section gone, every later section untouched. Pinned below.
* wordreport lost three breaks outright, because it uses each of its three
  flagged blocks exactly **once**: under a rule that exempts ``::1``, a block
  used once never breaks at all. That is what those templates did before the
  rule was briefly made uniform.
* **no stored reference moved.** A page break is written into an empty
  paragraph and the differential comparison reads text, so a case's
  ``expected/*.json`` cannot see one either way.

Why this module exists
----------------------
Until it was written the suite had **no** behavioural coverage of
``breakbefore``: ``tests/02_basetest/tag/test_tag.py`` pins that the argument
parses and that numbering a clone keeps it, and nothing asserted a page break
ever reached a document. The rule was carried by one line on a board and by
three shipped templates nobody had measured — which is how it stayed wrong,
and why the references above could not have caught the fix either.

The one path that never breaks
------------------------------
A *simple* tag cloned at a marker — ``<text:plain/>`` and its kind — goes
through ``DocParagraphElement.copy()``, which has no page-break step at all;
only the block classes inherit ``StructuredElement.copy()``. Unchanged by the
above, and out of scope here.
"""

from pathlib import Path
import shutil
import sys

import pytest

docx = pytest.importorskip('docx')
from docx.oxml.ns import qn                                    # noqa: E402

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = THIS_DIR.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *          # noqa: F401,F403  (brings reset_state, os)
from common_case import CaseConfig, run_docx_case              # noqa: E402

#: `04_examples/essay`, reached from here because what it demonstrates is a
#: property of *this* rule, not of that case: the essay test would have no
#: reason to explain the number it sees.
EXAMPLES = CASE_ROOT.parent.parent / '04_examples'

HEADER = ('_scriptum_:\n'
          '  version: 4\n'
          '  documenttype: docx\n'
          '  datadir: .\n'
          '_content_:\n')


def template(*, flag=' breakbefore', nested_flag=' breakbefore'):
    """``section:main`` holding ``alpha``, which holds ``one``.

    Both blueprints, each flagged independently so a test can put the
    argument on one level and not the other. Ordinary prose around every
    block, so a break's *position* is visible and not only its count.
    """
    document = docx.Document()
    for line in ('<section:main>MAIN',
                 'before alpha',
                 f'<subsection:alpha template{flag}>ALPHA <head/>',
                 'alpha: before one',
                 f'<subsubsection:one template{nested_flag}>ONE <item/>',
                 'one body',
                 '</subsubsection:one>',
                 'alpha: after one',
                 '</subsection:alpha>',
                 'after alpha'):
        document.add_paragraph(line)
    document.add_section()
    # A section ends with its break paragraph, and that paragraph must carry
    # the closing tag -- nothing may sit between the two.
    document.paragraphs[-1].text = 'main: end</section:main>'
    document.add_paragraph('<section:template>')
    document.add_paragraph('</section:template>')
    return document


def generate(tmp_path, body, built=None):
    """Build *body* against the template and return the finished path."""
    import Scriptum

    (built or template()).save(tmp_path / 'built.docx')
    (tmp_path / 'case.yaml').write_text(
        HEADER + '  - section:main:\n' + body, encoding='utf-8')

    os.chdir(tmp_path)
    rdf = Scriptum.ReportDataFile('case.yaml')
    assert not rdf.errors, rdf.errors
    managed = Scriptum.ManagedDocx('built.docx', rdf)
    managed.typesetting(rdf)
    managed.save('out.docx')
    return tmp_path / 'out.docx'


def broken(path):
    """The document as ``(breaks before it, text)`` per non-empty paragraph.

    A page break is written into a paragraph of its own, so what matters is
    how many such paragraphs stand between one piece of text and the next --
    which is exactly what the essay measurement below has to count.
    """
    lines, pending = [], 0
    for paragraph in docx.Document(path).paragraphs:
        if any(b.get(qn('w:type')) == 'page'
               for b in paragraph._p.iter(qn('w:br'))):
            pending += 1
        if paragraph.text.strip():
            lines.append((pending, paragraph.text.strip()))
            pending = 0
    return lines


def entries(name, *heads, level='subsection', indent=6):
    tag = {'subsection': 'head', 'subsubsection': 'item'}[level]
    pad = ' ' * indent
    return ''.join(f'{pad}- {level}:{name}:\n{pad}    - {tag}: {head}\n'
                   for head in heads)


# ----------------------------------------------------------- the rule today

def test_a_block_used_once_gets_no_break(tmp_path):
    """Instance 1 starts where the blueprint stood, so nothing is added.

    This is the whole of the change: a template that flags a block and then
    uses it once comes out with no page break at all -- which is what
    wordreport's three flagged blocks do, and what they did before the rule
    was briefly made uniform.
    """
    assert broken(generate(tmp_path, entries('alpha', 'A1'))) == [
        (0, 'MAIN'),
        (0, 'before alpha'),
        (0, 'ALPHA A1'),          # <- no break: it starts where it stands
        (0, 'alpha: before one'),
        (0, 'alpha: after one'),
        (0, 'after alpha'),
    ]


def test_every_instance_after_the_first_gets_one(tmp_path):
    """Three instances, two breaks -- in front of the second and the third,
    each standing between the instance before it and its own first
    paragraph."""
    said = broken(generate(tmp_path, entries('alpha', 'A1', 'A2', 'A3')))

    assert [(n, text) for n, text in said if text.startswith('ALPHA')] == \
        [(0, 'ALPHA A1'), (1, 'ALPHA A2'), (1, 'ALPHA A3')]
    assert sum(n for n, _ in said) == 2, 'no break anywhere else'


def test_without_the_argument_nothing_is_added(tmp_path):
    """The control. The same document against a template that does not carry
    the argument has no page breaks at all."""
    said = broken(generate(tmp_path, entries('alpha', 'A1', 'A2'),
                           built=template(flag='', nested_flag='')))

    assert sum(n for n, _ in said) == 0, said


def test_the_rule_counts_instances_per_level(tmp_path):
    """Each level counts its own instances. Repeating ``one`` inside the first
    ``alpha`` breaks before ``N2`` and not before ``N1``, and the same holds
    a level up -- both exercised in one document, so a rule that only looked
    at the outermost depth could not pass this."""
    body = (entries('alpha', 'A1')
            + entries('one', 'N1', 'N2', level='subsubsection', indent=10)
            + entries('alpha', 'A2'))
    said = broken(generate(tmp_path, body))

    assert said == [
        (0, 'MAIN'),
        (0, 'before alpha'),
        (0, 'ALPHA A1'),         # ::1 at depth 1
        (0, 'alpha: before one'),
        (0, 'ONE N1'),           # ::1 at depth 2
        (0, 'one body'),
        (1, 'ONE N2'),           # ::2 at depth 2
        (0, 'one body'),
        (0, 'alpha: after one'),
        (1, 'ALPHA A2'),         # ::2 at depth 1
        (0, 'alpha: before one'),
        (0, 'alpha: after one'),
        (0, 'after alpha'),
    ]


def test_the_outer_level_can_break_while_the_inner_does_not(tmp_path):
    """The argument is per block, not inherited: flagging ``alpha`` alone
    leaves ``one`` unbroken however often it repeats."""
    body = (entries('alpha', 'A1')
            + entries('one', 'N1', 'N2', level='subsubsection', indent=10)
            + entries('alpha', 'A2'))
    said = broken(generate(tmp_path, body, built=template(nested_flag='')))

    assert [(n, text) for n, text in said if text.startswith(('ALPHA', 'ONE'))] == \
        [(0, 'ALPHA A1'), (0, 'ONE N1'), (0, 'ONE N2'), (1, 'ALPHA A2')]


def test_an_unused_blueprint_leaves_no_stray_break(tmp_path):
    """A blueprint nobody names is pruned, and the page break goes with it --
    a break is added when a clone is made, never by the blueprint standing
    there. Worth pinning: the break paragraph is a separate paragraph, and a
    pruning pass that missed it would leave a blank page behind."""
    said = broken(generate(tmp_path, '      - text:description: nothing used\n'))

    assert sum(n for n, _ in said) == 0, said
    assert [text for _, text in said] == ['MAIN', 'before alpha', 'after alpha']


# ------------------------------------------- what a reversal would move

def test_the_shipped_essay_keeps_its_manual_pair_and_nothing_more(tmp_path):
    """The change, measured on the template that showed it was needed.

    ``04_examples/essay`` flags ``subsection:content``, and its template also
    carries **two manual page breaks** in front of that blueprint, put there
    before the rule was ever uniform. While ``::1`` broke as well, its first
    content section came out behind three breaks; it now comes out behind the
    author's two, and every later section behind the one the rule gives it.

    The measurement lives with the rule rather than with the essay case
    because it is the rule it is about -- and because no ``expected/*.json``
    can see it: a page break is an empty paragraph, and the differential
    comparison reads text.
    """
    essay = EXAMPLES / 'essay'
    if not (essay / 'essay.docx').is_file():
        pytest.skip(f'the essay case is not at {essay}')

    document = run_docx_case(CaseConfig(
        name='essay', case_dir=essay, document_name='essay.yaml',
        template_doc_name='essay.docx', output_name='breaks.docx',
        include_patterns=['essay*.yaml', 'essay.docx'],
        data_source_dir=essay / 'data', finish=False, createpdf=False), tmp_path)

    # which headings belong to subsection:content, read out of the document
    # rather than guessed: the essay's other blocks are Heading 1 too, and
    # picking the broken ones would be assuming the answer
    lines = (essay / 'essay.yaml').read_text(encoding='utf-8').splitlines()
    heads = [after.split('head:', 1)[1].strip().strip('\'"')
             for line, after in zip(lines, lines[1:])
             if line.strip() == '- subsection:content:' and 'head:' in after]
    assert len(heads) > 3, f'expected the essay to repeat its content block: {heads}'

    counted = {text: n for n, text in broken(document)}
    per_section = [counted[head] for head in heads]

    assert per_section[0] == 2, \
        f"the template's own two manual breaks, and nothing added: {per_section}"
    assert set(per_section[1:]) == {1}, \
        f'every later instance gets exactly one: {per_section}'
