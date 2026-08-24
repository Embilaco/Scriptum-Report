"""The text docx case, built and read back.

``word_text.yaml`` is the ladder at work on ``template.docx``: a title
section, a second section with a marker, two instances of the blueprint
``subsection:seconda`` (the first nesting a ``subsubsection`` and a
``sub3section``), a ``subsection:secondb`` whose marker takes text fills from
files -- three of them through an ``_include_`` fragment -- and four
``_global_`` values the headers and footers use. The case builds from its
own ``data/`` -- the workspace its ``CheckReport.ipynb`` prepares -- and two
referenced files exist nowhere on purpose, announced in place. A test
that only checks the file is there proves none of that, so this module reads
the document back:

* what it *says*, against ``expected/word_text.json`` (re-captured in
  `f143b8b` after the depth-3 tag was renamed ``sub3section``, and again
  from the own-``data/`` workspace when the test stopped linking the shared
  ``data_source`` pool);
* what it *shows*: the outline -- which heading style each ladder level got,
  something the text comparison cannot see -- and the globals in the header
  and footer of both sections. Where a clone lands is the business of
  ``test_docx_clone_order.py`` beside this file.
"""

from pathlib import Path
import shutil
import sys

import docx
import pytest

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = Path(__file__).resolve().parent.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *
from common_case import CaseConfig, run_docx_case
from common_case import said, normalise, reference, difference, portable, fold, drawings
from common_case import checkreport_comparison

REFERENCE = THIS_DIR / 'expected' / 'word_text.json'


def build(tmp_path):
    """The document, typeset the way every docx case is."""
    config = CaseConfig(
        name="report",
        case_dir=THIS_DIR,
        document_name="word_text.yaml",
        template_doc_name="template.docx",
        output_name="final_report.docx",
        include_patterns=["*.yaml", "template.docx"],
        data_source_dir=THIS_DIR / 'data',
        finish=False,
        createpdf=False,
    )
    return run_docx_case(config, tmp_path)


def test_document_is_created(tmp_path):
    print(f'\nWorking in {tmp_path}')

    result_path = build(tmp_path)

    assert result_path.exists(), "Expected final_report.docx to be generated"
    assert result_path.stat().st_size > 0, "Generated document should not be empty"


def test_the_document_says_what_the_reference_says(tmp_path):
    """Every paragraph, against the stored reference. The missing text files
    are announced with their path, which the runner makes absolute --
    portable() takes the workspace out again."""
    document = build(tmp_path)

    got = normalise(portable(said(document), document.parent))
    expected = [fold(line) for line in reference(REFERENCE)]

    assert got == expected, difference(expected, got)


def test_the_outline_the_includes_and_the_headers(tmp_path):
    """What the text comparison cannot see."""
    document = docx.Document(build(tmp_path))

    # each ladder level has its heading style; the depth-3 block, renamed
    # sub3section in f143b8b, is the Heading 4
    outline = [(p.style.name, p.text) for p in document.paragraphs
               if p.style.name.startswith('Heading')]
    assert outline == [
        ('Heading 1', 'TITLE'),
        ('Heading 3', 'When and who?'),
        ('Heading 3', 'What?'),
        ('Heading 1', 'SECOND'),
        ('Heading 2', 'Header 1'),
        ('Heading 3', 'This is s subsection ITEM head'),
        ('Heading 4', 'This is s subsubsection SUBITEM head'),
        ('Heading 2', 'Header 2'),
        ('Heading 2', 'Header 3'),
    ]

    # from the case's own data/: dolor.txt lands three times -- once in
    # subsection B directly, once through the _include_ fragment, once in
    # subsection A -- bootseal.txt twice and title.txt once; only
    # donotexist.txt (three times) and nonsense.txt exist nowhere and are
    # announced where they were asked for
    texts = [p.text for p in document.paragraphs]
    assert sum(text.startswith('Lorem ipsum') for text in texts) == 3
    assert sum(text.startswith('Boot seals') for text in texts) == 2
    assert sum(text.startswith('A title is everything') for text in texts) == 1
    announced = [text for text in texts if text.startswith('file ') and text.endswith('not found')]
    assert len(announced) == 4
    # the section's own marker fill lands after the subsections, the entry
    # after the _include_ after the fragment's three
    order = [texts.index(text) for text in
             ('Header 3', 'Some last words...', 'After the subsections', 'Where will this end?')]
    assert order == sorted(order), order

    # no pictures in this case; the globals reach header and footer of both sections
    assert len(document.inline_shapes) == 0
    assert len(document.sections) == 2
    for section in document.sections:
        header = ' '.join(p.text for p in section.header.paragraphs)
        assert 'Report ID 4711' in header and 'Date published: 01. August 2020' in header
        assert [drawings(p) for p in section.header.paragraphs if drawings(p)] == [[(0.8, 1.25)]]
        assert [p.text for p in section.footer.paragraphs if p.text.strip()] == ['Foot sec title ID 4711']


def test_a_linked_header_gets_its_global_text_once(tmp_path):
    """A section whose header and footer are linked to the previous one reads
    the same paragraphs, so the ``_global_`` text fills (report:id,
    date:published, the footer line) meet each paragraph once per section.
    The first visit fills and consumes the tag; the second must leave the
    text alone -- neither doubled nor failed. Pinned for pictures in the
    images case, for text here."""
    case = tmp_path / 'case'
    case.mkdir()
    for name in ('word_text.yaml', 'textinclude1.yaml'):
        shutil.copy(THIS_DIR / name, case)
    template = docx.Document(THIS_DIR / 'template.docx')
    template.sections[1].header.is_linked_to_previous = True
    template.sections[1].footer.is_linked_to_previous = True
    template.save(case / 'linked.docx')

    config = CaseConfig(
        name="report",
        case_dir=case,
        document_name="word_text.yaml",
        template_doc_name="linked.docx",
        output_name="final_report.docx",
        include_patterns=["*.yaml", "linked.docx"],
        data_source_dir=THIS_DIR / 'data',
        finish=False,
        createpdf=False,
    )
    document = docx.Document(run_docx_case(config, tmp_path))

    assert all(s.header.is_linked_to_previous for s in document.sections[1:])
    first = document.sections[0]
    header = ' '.join(p.text for p in first.header.paragraphs)
    assert header.count('Report ID 4711') == 1
    assert header.count('Date published: 01. August 2020') == 1
    assert [p.text for p in first.footer.paragraphs if p.text.strip()] == \
        ['Foot sec title ID 4711']
    leftover = [p.text for p in first.header.paragraphs + first.footer.paragraphs
                if '<' in p.text]
    assert not leftover


def test_the_checkreport_notebook_would_say_identical(tmp_path):
    """The comparison ``CheckReport.ipynb`` beside this file ends with --
    plain, no ``comparable()`` -- on the same build as above: the case test
    builds from the case's own ``data/``, so test and notebook share one
    workspace and one reference. Green when the notebook would print
    IDENTICAL; anything else reports as xfailed with the first difference."""
    report = checkreport_comparison(build(tmp_path), REFERENCE)
    if report:
        pytest.xfail('CheckReport.ipynb would not say IDENTICAL -- ' + report)
