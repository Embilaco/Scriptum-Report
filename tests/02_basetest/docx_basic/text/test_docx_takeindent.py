"""``takeindent`` on a marker: what is added there takes the marker's indent.

A block in ``<section:template>`` carries the indentation of the place it was
written, which is nowhere in particular — and everything added at a marker
therefore arrived at the template section's indent, whatever the marker's own
was. In a nested subsection that meant content falling out to the left margin
between two paragraphs that both sat well inside it.

``<marker:content takeindent/>`` is the template saying *content added here
takes my indentation*. The flag sits on the **marker**, not on the block,
because the marker is the donor: one blueprint is added at many markers of
many depths — ``text:generic`` is added at four in ``word_text.yaml`` alone,
at 0, 1.25, 2.50 and 3.75 cm — and only the marker knows which of them it is.

Two things this pins that no shipped fixture can see
----------------------------------------------------
* **The shift is relative, not absolute.** Every paragraph moves by the same
  delta — the marker's indent less the block's own first paragraph — so a
  block that indents internally keeps its shape. Every shipped block sits at
  zero, where relative and absolute agree exactly, so only a block built here
  can tell them apart.
* **The marker's indent may come from its style**, and python-docx reports
  direct formatting only: ``paragraph_format.left_indent`` is ``None`` both
  for *no indent* and for *inherited*. Reading it naively makes the flag a
  silent no-op on any template whose marker sits in an indented style — so
  the effective value is resolved through the style chain and the document
  defaults.
"""

from pathlib import Path
import sys

import pytest

docx = pytest.importorskip('docx')

from docx.shared import Cm

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = THIS_DIR.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *          # noqa: F401,F403  (brings reset_state, os)

HEADER = ('_scriptum_:\n'
          '  version: 4\n'
          '  documenttype: docx\n'
          '  datadir: .\n'
          '_content_:\n')


def template(markers, blockLines, indentedStyle=None):
    """A section of subsections, each holding one marker, plus one text block.

    *markers* is a list of ``(indent_cm_or_None, flagged)``, one subsection
    each. *blockLines* is a list of ``(text, indent_cm_or_None)`` making up
    ``<text:block>`` in the template section. ``indentedStyle``, when given,
    is the indent in cm of a paragraph **style** the flagged markers wear
    instead of carrying the indent directly.
    """
    document = docx.Document()

    if indentedStyle is not None:
        from docx.enum.style import WD_STYLE_TYPE
        style = document.styles.add_style('Marker', WD_STYLE_TYPE.PARAGRAPH)
        style.paragraph_format.left_indent = Cm(indentedStyle)

    document.add_paragraph('<section:main>MAIN')
    for number, (indent, flagged) in enumerate(markers, start=1):
        flag = ' takeindent' if flagged else ''
        # The whole subsection sits at the indent, marker included -- which is
        # the shape the complaint comes from: content added at the marker used
        # to land at the margin between two paragraphs well inside it.
        body = [document.add_paragraph(f'<subsection:sub{number}>SUB{number}'),
                document.add_paragraph(f'<marker:content{flag}/>'),
                document.add_paragraph(f'END SUB{number}</subsection:sub{number}>')]
        for paragraph in body:
            if indentedStyle is not None:
                paragraph.style = document.styles['Marker']
            if indent is not None:
                paragraph.paragraph_format.left_indent = Cm(indent)

    document.add_section()
    # A section ends with its break paragraph, and that paragraph must carry
    # the closing tag -- nothing may sit between the two.
    document.paragraphs[-1].text = '</section:main>'

    document.add_paragraph('<section:template>')
    document.add_paragraph('<text:block>')
    for text, indent in blockLines:
        paragraph = document.add_paragraph(text)
        if indent is not None:
            paragraph.paragraph_format.left_indent = Cm(indent)
    # An entry needs a value -- `- text:block:` alone reads as a container --
    # so the block carries a placeholder for the document to name. A scalar
    # would also stand, at the cost of the warning that it goes nowhere.
    document.add_paragraph('<placeholder:one/>')
    document.add_paragraph('</text:block>')
    document.add_paragraph('</section:template>')
    return document


def generate(tmp_path, markers, blockLines, indentedStyle=None):
    """Add ``text:block`` at every marker and read the indents back."""
    import Scriptum

    template(markers, blockLines, indentedStyle).save(tmp_path / 'built.docx')

    body = ''
    for number in range(1, len(markers) + 1):
        body += (f'      - subsection:sub{number}:\n'
                 f'          - marker:content:\n'
                 f'              - text:block:\n'
                 f'                  placeholder:one: filled\n')
    (tmp_path / 'case.yaml').write_text(
        HEADER + '  - section:main:\n' + body, encoding='utf-8')

    os.chdir(tmp_path)
    rdf = Scriptum.ReportDataFile('case.yaml')
    assert not rdf.errors, rdf.errors
    managed = Scriptum.ManagedDocx('built.docx', rdf)
    managed.typesetting(rdf)
    managed.save('out.docx')

    said = []
    for paragraph in docx.Document(tmp_path / 'out.docx').paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        indent = paragraph.paragraph_format.left_indent
        said.append((text, None if indent is None else round(indent.cm, 2)))
    return said


def indentOf(said, text):
    """The indent read back for the one line saying *text*."""
    hits = [indent for line, indent in said if line == text]
    assert len(hits) == 1, f'{text!r} appears {len(hits)} times in {said}'
    return hits[0]


BLOCK = [('opening line', None), ('closing line', None)]


# ------------------------------------------------------- what it does

def test_without_the_flag_the_block_keeps_the_templates_indent(tmp_path):
    """Today's behaviour, pinned: a plain marker donates nothing.

    The two lines land between an indented subsection's own paragraphs and
    do not join them -- which is the whole complaint, and the reason the
    flag exists. It stays the default: a template that says nothing gets
    what it has always got.
    """
    said = generate(tmp_path, [(3.0, False)], BLOCK)

    assert indentOf(said, 'END SUB1') == 3.0
    assert indentOf(said, 'opening line') is None
    assert indentOf(said, 'closing line') is None


def test_the_flag_moves_the_block_to_the_markers_indent(tmp_path):
    said = generate(tmp_path, [(3.0, True)], BLOCK)

    assert indentOf(said, 'END SUB1') == 3.0
    assert indentOf(said, 'opening line') == 3.0
    assert indentOf(said, 'closing line') == 3.0


def test_the_shift_is_relative_so_the_block_keeps_its_own_shape(tmp_path):
    """The block's middle line is indented 1 cm beyond its neighbours.

    Absolute would flatten the three to 3.0; relative moves all three by the
    same delta and the 1 cm survives. No shipped block indents internally,
    so this template is the only instrument that can tell the two apart.
    """
    said = generate(tmp_path, [(3.0, True)],
                    [('opening line', None),
                     ('middle line', 1.0),
                     ('closing line', None)])

    assert indentOf(said, 'opening line') == 3.0
    assert indentOf(said, 'middle line') == 4.0
    assert indentOf(said, 'closing line') == 3.0


def test_a_block_that_starts_indented_is_shifted_by_the_difference(tmp_path):
    """The delta is measured from the block's **first** paragraph.

    A block written at 0.5 cm and added at a 3.0 cm marker moves by 2.5, so
    its first line lands on the marker and everything keeps its offset.
    """
    said = generate(tmp_path, [(3.0, True)],
                    [('opening line', 0.5),
                     ('middle line', 1.5),
                     ('closing line', 0.5)])

    assert indentOf(said, 'opening line') == 3.0
    assert indentOf(said, 'middle line') == 4.0
    assert indentOf(said, 'closing line') == 3.0


def test_a_markers_indent_from_its_style_is_honoured(tmp_path):
    """The trap: ``paragraph_format.left_indent`` is None for an inherited
    indent, so reading it directly would make the flag do nothing at all.

    The marker here carries no direct indent -- only a style that has one.
    """
    said = generate(tmp_path, [(None, True)], BLOCK, indentedStyle=2.0)

    assert indentOf(said, 'opening line') == 2.0
    assert indentOf(said, 'closing line') == 2.0


def test_a_marker_at_the_margin_changes_nothing(tmp_path):
    """Flagged, but with nothing to donate: the delta is zero and no indent
    is written, so the block keeps whatever its own style gives it."""
    said = generate(tmp_path, [(None, True)], BLOCK)

    assert indentOf(said, 'opening line') is None
    assert indentOf(said, 'closing line') is None


def test_one_block_added_at_two_markers_lands_at_each(tmp_path):
    """Why the flag belongs on the marker and not on the block.

    One blueprint, two markers, two depths, two right answers -- which a
    flag on the block could not give.
    """
    said = generate(tmp_path, [(1.0, True), (4.0, True)], BLOCK)

    assert [indent for line, indent in said if line == 'opening line'] == [1.0, 4.0]
    assert [indent for line, indent in said if line == 'closing line'] == [1.0, 4.0]


def test_an_unflagged_marker_beside_a_flagged_one_is_untouched(tmp_path):
    """The flag is per marker, so the two do not interfere."""
    said = generate(tmp_path, [(1.0, False), (4.0, True)], BLOCK)

    assert [indent for line, indent in said if line == 'opening line'] == [None, 4.0]


def test_a_table_moves_with_its_own_caption(tmp_path):
    """A table block is a table *and* the paragraphs around it.

    Paragraphs alone would move the caption and leave the table at the
    margin, which is a worse document than the one the flag was added to
    fix. A table indents through ``w:tblInd``, which python-docx does not
    expose, so this is the one place the back end writes the XML itself.
    """
    import Scriptum
    from docx.oxml.ns import qn
    from docx.shared import Twips

    document = docx.Document()
    document.add_paragraph('<section:main>MAIN')
    marker = document.add_paragraph('<marker:content takeindent/>')
    marker.paragraph_format.left_indent = Cm(3.0)
    document.add_section()
    document.paragraphs[-1].text = '</section:main>'
    document.add_paragraph('<section:template>')
    document.add_paragraph('<table:narrow>')
    document.add_table(rows=2, cols=2)
    document.add_paragraph('Table: <description/>')
    document.add_paragraph('</table:narrow>')
    document.add_paragraph('</section:template>')
    document.save(tmp_path / 'built.docx')

    (tmp_path / 'rows.csv').write_text('a,b\n1,2\n', encoding='utf-8')
    (tmp_path / 'case.yaml').write_text(
        HEADER + '  - section:main:\n'
                 '      - marker:content:\n'
                 '          - table:narrow: {file: rows.csv, description: a caption}\n',
        encoding='utf-8')

    os.chdir(tmp_path)
    rdf = Scriptum.ReportDataFile('case.yaml')
    assert not rdf.errors, rdf.errors
    managed = Scriptum.ManagedDocx('built.docx', rdf)
    managed.typesetting(rdf)
    managed.save('out.docx')

    out = docx.Document(tmp_path / 'out.docx')

    caption = [p for p in out.paragraphs if p.text.strip().startswith('Table:')]
    assert len(caption) == 1, [p.text for p in out.paragraphs]
    assert round(caption[0].paragraph_format.left_indent.cm, 2) == 3.0

    assert len(out.tables) == 1
    indent = out.tables[0]._tbl.tblPr.find(qn('w:tblInd'))
    assert indent is not None, 'the table stayed at the margin'
    assert round(Twips(int(indent.get(qn('w:w')))).cm, 2) == 3.0
