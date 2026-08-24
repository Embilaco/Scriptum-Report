"""``findTemplate`` on a name that exists twice: section:template wins, loudly.

`8c3c210` widened the bare-name lookup -- what an ``add`` at a marker carries
-- to every block flagged ``template``; the shipped templates keep blueprint
names unique, so the collision rules ran on trust: the template section
outranks an in-content blueprint, any multi-hit is warned about (names being
documented as unique), and the loser is not touched -- an unused in-content
blueprint **ships**, by the same decision.

The document uses the colliding name the way the tables fixture uses
``table:orange``: instance 1 fills the in-content blueprint in place,
instance 2 is added at the marker and resolved by bare name. (Using the name
*only* through the marker is a different, unsettled story: the lookup takes
the collision winner but the fill follows the address to the in-content
namesake -- filed as a question on the DOCX board, not pinned here.)

Templates are built in-test, the way ``test_docx_clone_order.py`` builds its
own; like there, each carries a ``section:template`` -- ``typesetting``
deletes that section by name and does not expect a document without one. The
winner is read off the finished document: a CSV fill lands below the sample
row, whose second cell carries each blueprint's fingerprint.
"""

from pathlib import Path
import os
import sys

import docx

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = THIS_DIR.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *          # noqa: F401,F403

HEADER = ('_scriptum_:\n'
          '  version: 4\n'
          '  documenttype: docx\n'
          '  datadir: ./data\n'
          '_content_:\n')

#: Instance 1 in place, instance 2 added at the marker by bare name.
BODY = ('  - section:mix:\n'
        '      - table:dup: {file: mini1.csv}\n'
        '      - marker:content:\n'
        '          - table:dup: {file: mini2.csv}\n')


def blueprint(document, fingerprint, *, flagged):
    """A ``table:dup`` blueprint block whose second cell says *fingerprint*.

    Two rows, like every shipped blueprint table: the first is the sample
    row the fingerprint rides in, the second is where a CSV fill lands."""
    flag = ' template' if flagged else ''
    document.add_paragraph(f'<table:dup{flag}>')
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 1).text = fingerprint
    document.add_paragraph('</table:dup>')


def collision_template(*, in_template_section):
    """``table:dup`` twice: in ``section:mix`` (*FromContent*), and either in
    ``section:template`` (*FromTemplateSection*) or flagged in a second
    content section (*FromSecond*, with an empty template section after)."""
    document = docx.Document()
    document.add_paragraph('<section:mix>MIX')
    document.add_paragraph('<marker:content/>')
    blueprint(document, 'FromContent', flagged=True)
    document.add_section()
    # A section ends with its break paragraph, and that paragraph must carry
    # the closing tag -- nothing may sit between the two.
    document.paragraphs[-1].text = '</section:mix>'
    if in_template_section:
        document.add_paragraph('<section:template>')
        blueprint(document, 'FromTemplateSection', flagged=False)
        document.add_paragraph('</section:template>')
    else:
        document.add_paragraph('<section:extra>EXTRA')
        blueprint(document, 'FromSecond', flagged=True)
        document.add_section()
        document.paragraphs[-1].text = '</section:extra>'
        document.add_paragraph('<section:template>')
        document.add_paragraph('</section:template>')
    return document


def generate(tmp_path, template):
    """Build the collision document and return its tables."""
    import Scriptum

    template.save(tmp_path / 'built.docx')
    (tmp_path / 'data').mkdir(exist_ok=True)
    (tmp_path / 'data' / 'mini1.csv').write_text('ZZinplace\n', encoding='utf-8')
    (tmp_path / 'data' / 'mini2.csv').write_text('ZZadded\n', encoding='utf-8')
    (tmp_path / 'case.yaml').write_text(HEADER + BODY, encoding='utf-8')

    os.chdir(tmp_path)
    rdf = Scriptum.ReportDataFile('case.yaml')
    document = Scriptum.ManagedDocx('built.docx', rdf)
    document.typesetting(rdf)
    document.save('out.docx')

    return docx.Document('out.docx').tables


def fingerprints(tables):
    return [table.cell(0, 1).text for table in tables]


def cells(table):
    return [cell.text for row in table.rows for cell in row.cells]


def test_the_template_section_wins_the_collision(tmp_path, capsys):
    """The add clones the ``section:template`` blueprint, not the in-content
    one -- and says so: a multi-hit is warned about even when the ranking
    settles it. Instance 1 fills the in-content blueprint in place, which
    then ships; the template section itself is removed as always."""
    tables = generate(tmp_path, collision_template(in_template_section=True))
    out = capsys.readouterr().out

    assert fingerprints(tables) == ['FromTemplateSection', 'FromContent'], \
        'the clone at the marker first, then the in-place-filled blueprint'
    assert 'ZZadded' in cells(tables[0]), 'the marker add got its own CSV'
    assert 'ZZinplace' in cells(tables[1]), 'instance 1 filled in place'
    assert "template 'table:dup' is ambiguous" in out
    assert 'section:template' in out.split('is ambiguous', 1)[1].splitlines()[0], \
        'the warning lists the hits, the template section first'


def test_without_a_template_section_hit_the_first_flagged_block_wins(tmp_path, capsys):
    """Both blueprints live in content sections: the first in document order
    is cloned to the marker, the ambiguity is still warned about, and the
    unused loser ships on unchanged."""
    tables = generate(tmp_path, collision_template(in_template_section=False))
    out = capsys.readouterr().out

    assert fingerprints(tables) == ['FromContent', 'FromContent', 'FromSecond'], \
        'the clone of the first flagged block, then the two shipping blueprints'
    assert 'ZZadded' in cells(tables[0])
    assert 'ZZinplace' in cells(tables[1])
    assert not any(text.startswith('ZZ') for text in cells(tables[2])), \
        'the losing blueprint is untouched'
    assert "template 'table:dup' is ambiguous" in out
