"""Where a clone lands, and what is left of a blueprint afterwards.

A block whose tag says ``template`` is a **blueprint**: every instance of it
is a clone, the first included, and the blueprint itself is pruned at the
end. A block without the argument is content: its first instance is the block
itself and later instances clone it. Word thereby agrees with PowerPoint,
which always copies.

Placement follows from that. The first instance of a blueprint goes exactly
where the blueprint stands -- before its opening paragraph, which is then
removed with the rest of it -- so the finished document keeps the template's
order whatever order the data names things in. Further instances go before
the first *unclaimed* sibling, and claiming a child claims everything ahead of
it (``claimSubAnchor``): an earlier blueprint the data never used must not
stay at the front, or the next clone goes in upstream of the block it follows,
which is the defect reproduced here on two of the shipped templates.

``template_text.docx`` is the right template for most of this: ``section:
second`` holds ``<subsection:seconda template>`` -- with a nested
``<subsubsection:secondsub1 template>`` whose body says *Secondsub1* -- and
then ``<subsection:secondb template>``. The one case it cannot express, a
blueprint beside an ordinary sibling, gets a template built here.
"""

from pathlib import Path
import shutil
import sys

import docx

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = THIS_DIR.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *          # noqa: F401,F403  (brings reset_state)

TEMPLATE = 'template_text.docx'
MARKS = ('ZZintro', 'ZZsibling', 'ZZfirst', 'ZZsecond')

HEADER = ('_scriptum_:\n'
          '  version: 4\n'
          '  documenttype: docx\n'
          '  datadir: ./data\n'
          '_content_:\n')


def generate(tmp_path, second, template=TEMPLATE, section='section:second'):
    """Build a document from one section's body and return its paragraphs."""
    import Scriptum

    if isinstance(template, str):
        shutil.copy(THIS_DIR / template, tmp_path)
    else:
        template.save(tmp_path / 'built.docx')
        template = 'built.docx'
    (tmp_path / 'data').mkdir(exist_ok=True)

    (tmp_path / 'case.yaml').write_text(
        HEADER + f'  - {section}:\n' + second, encoding='utf-8')

    os.chdir(tmp_path)
    rdf = Scriptum.ReportDataFile('case.yaml')
    document = Scriptum.ManagedDocx(template, rdf)
    document.typesetting(rdf)
    document.save('out.docx')

    return [p.text.strip() for p in docx.Document('out.docx').paragraphs]


def marks(paragraphs):
    return [p for p in paragraphs if p in MARKS]


TWO_SECONDB = ('      - text:description: ZZintro\n'
               '      - subsection:secondb:\n'
               '          - head: ZZfirst\n'
               '      - subsection:secondb:\n'
               '          - head: ZZsecond\n')

SIBLING = ('      - subsection:seconda:\n'
           '          - head: ZZsibling\n')


# ------------------------------------------------------------ placement

def test_a_clone_follows_its_instance_when_an_earlier_blueprint_is_unused(tmp_path):
    """The regression. Before the fix this came back as intro, second, first."""
    assert marks(generate(tmp_path, TWO_SECONDB)) == \
        ['ZZintro', 'ZZfirst', 'ZZsecond']


def test_a_clone_still_follows_its_instance_when_the_blueprint_is_used(tmp_path):
    """The case that always worked, kept so the fix cannot break it."""
    body = '      - text:description: ZZintro\n' + SIBLING + TWO_SECONDB[
        len('      - text:description: ZZintro\n'):]
    assert marks(generate(tmp_path, body)) == \
        ['ZZintro', 'ZZsibling', 'ZZfirst', 'ZZsecond']


def test_first_instances_keep_the_template_order_whatever_the_data_order(tmp_path):
    """``secondb`` named before ``seconda`` still comes out after it.

    The first instance of a blueprint lands where the blueprint stands, so
    the data's order cannot reorder the template. A cursor that popped "the
    next free slot" regardless of which block was named would put secondb's
    clone into seconda's place here -- the design that was tried first.
    """
    body = ('      - subsection:secondb:\n'
            '          - head: ZZfirst\n'
            + SIBLING)
    assert marks(generate(tmp_path, body)) == ['ZZsibling', 'ZZfirst']


def test_claiming_a_child_claims_everything_ahead_of_it():
    """The rule on its own, without a document.

    A child that is no longer listed is not an error: a document may use blocks
    in an order the template does not hold them in, and the first such use
    already dropped the ones ahead of it.
    """
    from Scriptum._docx.structure import StructuredElement

    holder = StructuredElement.__new__(StructuredElement)
    holder.subAnchors = ['a', 'b', 'c', 'd']

    holder.claimSubAnchor('c')
    assert holder.subAnchors == ['d']

    holder.claimSubAnchor('a')          # already gone, and not an error
    assert holder.subAnchors == ['d']

    holder.claimSubAnchor('d')
    assert holder.subAnchors == []


# -------------------------------------------------------------- pruning

def test_an_unused_blueprint_leaves_nothing_behind(tmp_path):
    """Only ``secondb`` is used, so ``seconda`` -- and the nested blueprint
    inside it -- must not show up: its tags cleaned and its text intact was
    what the finished document used to carry."""
    said = generate(tmp_path, TWO_SECONDB)

    assert 'Secondsub1' not in said
    assert 'SubSub I' not in said
    assert 'Between the subsections' in said, 'ordinary content stays'


def test_a_blueprint_carried_inside_a_clone_is_pruned_when_unused(tmp_path):
    """A clone of ``seconda`` carries a copy of the ``secondsub1`` blueprint.
    The first instance uses it, the second does not -- so its text appears
    once, not twice. Twice is what a clone used to leak."""
    body = ('      - subsection:seconda:\n'
            '          - head: ZZfirst\n'
            '          - subsubsection:secondsub1:\n'
            '              - item: used here\n'
            '      - subsection:seconda:\n'
            '          - head: ZZsecond\n')
    said = generate(tmp_path, body)

    assert said.count('Secondsub1') == 1
    assert marks(said) == ['ZZfirst', 'ZZsecond']


# --------------------------------------- a blueprint beside ordinary content

def built_template():
    """``section:mix`` holds a blueprint, an ordinary subsection, a blueprint.

    No shipped template has the mixture, and it is the shape that separates
    "the first instance goes where its blueprint stands" from any rule that
    counts slots: name the three in reverse and the output must still read
    alpha, plain, omega.
    """
    document = docx.Document()
    for line in ('<section:mix>MIX',
                 '<subsection:alpha template><head/>',
                 '</subsection:alpha>',
                 '<subsection:plain><head/>',
                 '</subsection:plain>',
                 '<subsection:omega template><head/>',
                 '</subsection:omega>'):
        document.add_paragraph(line)
    document.add_section()
    # A section ends with its break paragraph, and that paragraph must carry
    # the closing tag -- nothing may sit between the two.
    document.paragraphs[-1].text = '</section:mix>'
    document.add_paragraph('<section:template>')
    document.add_paragraph('</section:template>')
    return document


def test_a_blueprint_beside_ordinary_content_keeps_its_place(tmp_path):
    body = ('      - subsection:omega:\n'
            '          - head: ZZomega\n'
            '      - subsection:plain:\n'
            '          - head: ZZplain\n'
            '      - subsection:alpha:\n'
            '          - head: ZZalpha\n')
    said = generate(tmp_path, body, template=built_template(),
                    section='section:mix')

    assert [p for p in said if p.startswith('ZZ')] == \
        ['ZZalpha', 'ZZplain', 'ZZomega']
