"""The simple docx case, built and read back.

``word_simple.yaml`` fills only the title section of ``template.docx`` -- a
name, a designer, a reason, the main model picture into a table cell, a text
file, ``date: now`` -- and four ``_global_`` values (report id, publication
date, version, status) that the template uses in the title table and in the
header of **every** section. The rest of the template ships unfilled: its
tags are cleaned, its blueprints and the template section pruned. A test that
only checks the file is there proves none of that, so this module reads the
document back:

* what it *says*, against ``expected/word_simple.json`` (captured at
  `44267a8` from the ``.rdf`` this fixture was translated from, by the
  differential harness this case has graduated from);
* what it *shows*: the picture in the title table at the tag's 8 cm, the
  template's own logo left alone in each header, and the tables' shapes.
"""

from pathlib import Path
import sys

import docx

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = Path(__file__).resolve().parent.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *
from common_case import CaseConfig, run_docx_case
from common_case import said, normalise, reference, difference, portable, fold, drawings

REFERENCE = THIS_DIR / 'expected' / 'word_simple.json'


def build(tmp_path):
    """The document, typeset the way every docx case is."""
    config = CaseConfig(
        name="report",
        case_dir=THIS_DIR,
        document_name="word_simple.yaml",
        template_doc_name="template.docx",
        output_name="final_report.docx",
        include_patterns=["*.yaml", "template.docx"],
        data_source_dir=DATA_SOURCE,
        finish=False,
        createpdf=False,
    )
    return run_docx_case(config, tmp_path)


def test_document_is_created(tmp_path):
    print(f'\nWorking in {tmp_path}')

    result_path = build(tmp_path)

    assert result_path.exists(), "Expected final_report.docx to be generated"
    assert result_path.stat().st_size > 0, "Generated document should not be empty"


def test_the_document_says_what_the_reference_says(tmp_path):
    """Every paragraph, then every table cell, against the stored reference --
    digits and weekday names collapsed on both sides (``date: now`` is per
    run, the reference is from another day), paths made portable."""
    document = build(tmp_path)

    got = normalise(portable(said(document), document.parent))
    expected = [fold(line) for line in reference(REFERENCE)]

    assert got == expected, difference(expected, got)


def test_the_picture_the_tables_and_the_headers(tmp_path):
    """What the text comparison cannot see."""
    document = docx.Document(build(tmp_path))

    # image:mainmodel (screw.png, square) sits in the first cell of the title
    # table, behind <Image:MainModel width=8cm/>: one picture, 8 cm wide
    title_table = document.tables[0]
    cell_pictures = [size for paragraph in title_table.cell(0, 0).paragraphs
                     for size in drawings(paragraph)]
    assert cell_pictures == [(8.0, 8.0)]
    assert 'Date: 25. Nov 2025' in title_table.cell(0, 1).text, 'the global date:published'

    # the template's own title-page picture is the only other inline one
    assert len(document.inline_shapes) == 1
    assert [drawings(p) for p in document.paragraphs if drawings(p)] == [[(4.61, 7.22)]]

    # the revisions table (fixed), two blueprint tables (unfilled, cleaned)
    assert [(len(t.rows), len(t.columns)) for t in document.tables] == [
        (1, 2), (5, 3), (2, 3), (2, 3)]
    assert [cell.text for cell in document.tables[1].rows[0].cells] == ['Date', 'Who', 'What']

    # the global report:id reached the header of every section, next to the
    # template's logo, which is left as it is
    assert len(document.sections) == 6
    for section in document.sections:
        texts = [p.text for p in section.header.paragraphs if p.text.strip()]
        assert any('Report Test a Screw' in text for text in texts), texts
        assert [drawings(p) for p in section.header.paragraphs if drawings(p)] == [[(0.8, 1.25)]]
        assert [p.text for p in section.footer.paragraphs if p.text.strip()] == ['pg. 2']

    assert document.core_properties.author.startswith('Scriptum ')


def test_the_documents_title_is_the_settings_default(tmp_path):
    """``word_simple.yaml`` sets no ``documenttitle``, so ``setproperties``
    writes the settings default -- the pptx twin pins the set case."""
    document = docx.Document(build(tmp_path))

    assert document.core_properties.title == 'Autoreport'
