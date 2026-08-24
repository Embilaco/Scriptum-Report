"""Umlauts, accents and emoji pass through a Word fill unchanged.

A value travels document -> loader -> task -> python-docx run, and nothing on
the way may re-encode it. No case fixture pinned that: they are plain ASCII
throughout, and Word's own smart-quote mangling is documented elsewhere. The
document title takes the same trip into the core properties.

The text-**file** route is pinned separately below, because it is broken in a
platform-dependent way: ``text_value.content`` opens the file without an
``encoding``, so the platform decides -- on Windows that is the ANSI codepage,
and UTF-8 content mojibakes or fails to decode (the CSV reader beside it says
``encoding='utf-8'`` explicitly). Reported as xfail wherever it happens.
"""

from pathlib import Path
import shutil
import sys

import docx
import pytest

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = THIS_DIR.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *          # noqa: F401,F403

DOCUMENT = """_scriptum_:
  version: 4
  documenttype: docx
  datadir: ./data
  documenttitle: 'Prüfbericht — Größenprüfung'
_content_:
  - section:title:
      - report:product_name: 'Größenprüfung äöü ß'
      - report:designer_name: 'Ægir Ångström 🚀'
{extra}"""

FILE_TEXT = 'Résumé ✓ 100 € — naïve façade'


def build(tmp_path, extra=''):
    import Scriptum

    shutil.copy(THIS_DIR / 'template.docx', tmp_path)
    (tmp_path / 'data').mkdir()
    (tmp_path / 'data' / 'unicode.txt').write_text(FILE_TEXT, encoding='utf-8')
    (tmp_path / 'case.yaml').write_text(DOCUMENT.format(extra=extra),
                                        encoding='utf-8')
    os.chdir(tmp_path)

    rdf = Scriptum.ReportDataFile('case.yaml')
    document = Scriptum.ManagedDocx('template.docx', rdf)
    document.typesetting(rdf)
    document.save('out.docx')
    return docx.Document('out.docx')


def said(document):
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        lines.extend(cell.text for row in table.rows for cell in row.cells)
    return lines


def test_nonascii_values_reach_the_document_verbatim(tmp_path):
    finished = build(tmp_path)
    texts = said(finished)

    assert any('Größenprüfung äöü ß' in text for text in texts), texts
    assert any('Ægir Ångström 🚀' in text for text in texts), texts
    assert finished.core_properties.title == 'Prüfbericht — Größenprüfung'


def test_a_nonascii_text_file_reaches_the_document(tmp_path):
    """Green where the platform's default encoding is UTF-8; xfailed where it
    is not (Windows ANSI), because ``text_value.content`` leaves the encoding
    to the platform. The fill degrades to nothing here rather than crashing,
    which is exactly why it needs a mark to be seen at all."""
    finished = build(tmp_path,
                     '      - text:description: {file: unicode.txt}\n')
    texts = said(finished)

    if not any(FILE_TEXT in text for text in texts):
        pytest.xfail("text files are read in the platform encoding "
                     "(text_value.content opens without encoding=) -- "
                     f"{FILE_TEXT!r} did not arrive")
