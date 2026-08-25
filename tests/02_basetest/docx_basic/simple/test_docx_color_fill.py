"""A ``color`` modifier paints the font of the text its fill writes.

The first thing a colour actually *does* in a document: the loader has long
parsed and diagnosed colours while both back ends dropped them. Now the run
that receives the fill's text is painted -- and only that run. Word splits a
tag into runs of its own, so template text next to the tag keeps its own
colour: the ``Report `` prefix beside the ``report:ID`` fill stays unpainted
below, which is the exactness this pins. The pptx twin is
``pptx-basic/simple/test_pptx_color_fill.py``.

Three routes, one colour notation each: a direct fill (named colour), a
``_global_`` fill (corporate hex), and a text-file fill (``rgb(...)``).
"""

from pathlib import Path
import shutil
import sys

import docx
from docx.shared import RGBColor

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = THIS_DIR.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *          # noqa: F401,F403

DOCUMENT = """_scriptum_:
  version: 4
  documenttype: docx
  datadir: ./data
_global_:
  report:id:
    text: Corporate Screw
    color: 'B00020'
_content_:
  - section:title:
      - report:product_name:
          text: A Crimson Screw
          color: crimson
      - text:description:
          file: note.txt
          color: rgb(0, 85, 150)
"""

NOTE = 'A blue corporate note'


def build(tmp_path, monkeypatch):
    import Scriptum

    shutil.copy(THIS_DIR / 'template.docx', tmp_path)
    (tmp_path / 'data').mkdir()
    (tmp_path / 'data' / 'note.txt').write_text(NOTE, encoding='utf-8')
    (tmp_path / 'case.yaml').write_text(DOCUMENT, encoding='utf-8')
    monkeypatch.chdir(tmp_path)

    rdf = Scriptum.ReportDataFile('case.yaml')
    document = Scriptum.ManagedDocx('template.docx')
    document.typesetting(rdf)
    document.save('out.docx')
    return docx.Document('out.docx')


def all_paragraphs(document):
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def colored(document):
    return {run.text: run.font.color.rgb
            for paragraph in all_paragraphs(document)
            for run in paragraph.runs if run.text}


def test_the_three_fill_routes_paint_their_runs(tmp_path, monkeypatch):
    colors = colored(build(tmp_path, monkeypatch))

    assert colors['A Crimson Screw'] == RGBColor(0xDC, 0x14, 0x3C)  # crimson
    assert colors['Corporate Screw'] == RGBColor(0xB0, 0x00, 0x20)  # global
    assert colors[NOTE] == RGBColor(0x00, 0x55, 0x96)               # rgb(...)


def test_template_text_beside_the_fill_keeps_its_colour(tmp_path, monkeypatch):
    """The template says ``Report <report:ID/>`` in one paragraph; the paint
    must land on the fill's run and leave the ``Report `` run alone."""
    document = build(tmp_path, monkeypatch)

    lines = [p for p in all_paragraphs(document)
             if p.text == 'Report Corporate Screw']
    assert lines, 'the report:ID paragraph is gone from the title section'
    runs = {run.text: run.font.color.rgb for run in lines[0].runs if run.text}

    assert runs['Report '] is None
    assert runs['Corporate Screw'] == RGBColor(0xB0, 0x00, 0x20)
