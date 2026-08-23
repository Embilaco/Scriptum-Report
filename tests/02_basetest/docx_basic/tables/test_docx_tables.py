"""The tables docx case, built and read back.

``word_tables.yaml`` fills the title section of ``template_table.docx`` with
tables from the case's own ``data/``: a fixed revisions table that takes the
``_global_`` author, ``table:inline`` from a CSV, a ``table:generic`` clone
added at the marker, ``table:orange`` -- a blueprint that lives in the title
section itself -- filled in place with its description ``from: row1``, and,
the case's deliberate failure path, a second ``table:orange`` *added* at a
marker: that blueprint is not in ``section:template``, and ``add`` looks only
there, so the back end refuses with two warnings and the second table never
appears. Whether it should is an open point on the DOCX board (*Can a block
flagged `template` outside section:template be added at a marker?*); until it
is decided the warnings are pinned here, so a back end that starts placing
that table announces itself. A test that only checks the file is there proves
none of that, so this module reads the document back:

* what it *says*, against ``expected/word_tables.json`` (re-captured at
  `a58702e`, when the fixture got its own data and real CSV files);
* what it *shows*: the tables' shapes and the CSV content in their cells.
"""

import re
from pathlib import Path
import sys

import docx

THIS_DIR = Path(__file__).resolve().parent
CASE_ROOT = Path(__file__).resolve().parent.parent
if str(CASE_ROOT) not in sys.path:
    sys.path.append(str(CASE_ROOT))

from _setup_docx_basic import *
from common_case import CaseConfig, run_docx_case
from common_case import said, normalise, reference, difference, portable, fold

REFERENCE = THIS_DIR / 'expected' / 'word_tables.json'

#: What this fixture makes the back end say: table:orange is a blueprint of
#: the title section, so instance 2 -- an add at marker:content -- finds no
#: template to clone and no exact block to fill, and is dropped.
EXPECTED_WARNINGS = [
    "WARNING: No such template in document: ['section:template', 'table:orange']",
    "WARNING: No exact match: ['section:title::1', 'table:orange::2']",
]


def build(tmp_path):
    """The document, typeset the way every docx case is."""
    config = CaseConfig(
        name="report",
        case_dir=THIS_DIR,
        document_name="word_tables.yaml",
        template_doc_name="template_table.docx",
        output_name="final_report.docx",
        include_patterns=["*.yaml", "template_table.docx"],
        data_source_dir=THIS_DIR / 'data',
        finish=False,
        createpdf=False,
    )
    return run_docx_case(config, tmp_path)


def test_document_is_created(tmp_path):
    print(f'\nWorking in {tmp_path}')

    result_path = build(tmp_path)

    assert result_path.exists(), "Expected final_report.docx to be generated"
    assert result_path.stat().st_size > 0, "Generated document should not be empty"


def test_the_document_says_what_the_reference_says(tmp_path):
    """Every paragraph, then every table cell, against the stored reference."""
    document = build(tmp_path)

    got = normalise(portable(said(document), document.parent))
    expected = [fold(line) for line in reference(REFERENCE)]

    assert got == expected, difference(expected, got)


def test_the_tables_their_content_and_the_refused_add(tmp_path, capsys):
    """What the text comparison cannot see, and what the run says about it."""
    document = docx.Document(build(tmp_path))
    warnings = [line for line in capsys.readouterr().out.splitlines() if 'WARNING' in line]

    assert warnings == EXPECTED_WARNINGS

    # title table, revisions table, table:inline from instructionsrocket.csv,
    # table:orange in place from table2.csv, the table:generic clone from
    # technologies.csv -- and no second table:orange
    tables = document.tables
    assert [(len(t.rows), len(t.columns)) for t in tables] == [
        (1, 2), (5, 3), (4, 3), (8, 3), (6, 4)]

    # date:creation is `date: now` in the default ISO datetime format
    assert re.search(r'^Date: \d{4}-\d{2}-\d{2} \d{2}:\d{2}', tables[0].cell(0, 1).text)

    # the fixed revisions table took the _global_ author
    assert [c.text for c in tables[1].rows[0].cells] == ['Date', 'Who', 'What']
    assert [c.text for c in tables[1].rows[1].cells] == ['', 'James Bond', 'Initial revision']

    # the CSV content, header rows and a data row each
    assert [c.text for c in tables[2].rows[0].cells] == ['Check', 'Task', 'Importance']
    assert [c.text for c in tables[2].rows[1].cells][:2] == ['Rocket', 'ask someone to get one for you']
    assert [c.text for c in tables[3].rows[0].cells] == ['Rank', 'Country', 'Income']
    assert [c.text for c in tables[3].rows[1].cells] == ['1', 'Luxembourg', '$140,941']
    assert [c.text for c in tables[4].rows[1].cells] == [
        'Technology', 'Typical Period', 'Key Strengths', 'Key Limitations']

    # the captions: the descriptions given, and the one taken from row1
    captions = [p.text for p in document.paragraphs if p.text.startswith('Table ')]
    assert captions == ['Table 1: rocket preparation', 'Table 1: Income by country', "Table 4: tech isn't it"]
    assert not any('non existing file' in p.text for p in document.paragraphs)
